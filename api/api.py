from flask import Flask, request, jsonify, send_file, make_response, send_from_directory
import openpyxl
import xlsxwriter
import numbers_parser
import pandas as pd
import re
import os

from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from flask_cors import CORS, cross_origin
from numbers_parser import Document
from openpyxl import load_workbook
from docx import Document as WordDocument

app = Flask(__name__, static_folder="../build", static_url_path='')

CORS(app) # enable CORS for all routes and origins

# Path to the generated Excel file
processed_file_path = "COSTCO.xlsx"

def get_column_data(rows, header_row, column_name):
    column_index = None
    data = []
    for index, cell in enumerate(header_row):
        if cell.value == column_name:
            column_index = index
            break
    if column_index is not None:
        for row in rows[1:]:
            cell = row[column_index]
            if cell.value is not None:
                data.append(cell.value)
    return data

def get_csv_column_info(df, column_name):
    if column_name in df.columns:
        return df[column_name].dropna().tolist()
    return []

def convert_csv_data(file_name):
    # Load the CSV file into a DataFrame
    df = pd.read_csv(file_name)
    
    # Get data from the DataFrame
    delivery_dates.extend(get_csv_column_info(df, "Requested Delivery Date"))
    ship_to_name.extend(get_csv_column_info(df, "Ship To Name"))
    ship_to_address.extend(get_csv_column_info(df, "Ship To Address 1"))
    ship_to_city.extend(get_csv_column_info(df, "Ship To City"))
    ship_to_state.extend(get_csv_column_info(df, "Ship To State"))
    ship_to_zip.extend(get_csv_column_info(df, "Ship to Zip"))
    ship_to_contact.extend(get_csv_column_info(df, "Ship To Contact"))
    sell_units.extend(get_csv_column_info(df, "Qty Ordered"))
    sku_num.extend(get_csv_column_info(df,"Buyers Catalog or Stock Keeping #"))
    
    # PO Numbers
    po_column = df['PO Number'] # Assuming the column name is "PO Number"
    unique_po_nums = set()  # Initialize a set to store unique PO Numbers
    # Iterate through the PO Number column and extend po_nums, skipping duplicates
    for po in po_column:
        if pd.notna(po) and po not in unique_po_nums:  # Check if the value is not NaN and not a duplicate
            po_nums.append(po)
            unique_po_nums.add(po)
    
    # Calculate 'CARDS PER UNIT' from 'Product/Item Description'
    if "Product/Item Description" in df.columns:
        df['CARDS PER UNIT'] = df['Product/Item Description'].str.extract(r'(\d+)/\$').astype(float)
        cards_per_units.extend(df['CARDS PER UNIT'].dropna().astype(int).tolist())
        df['PRICE'] = df['Product/Item Description'].str.extract(r'\$(\d+(?:\.\d+)?)').astype(float)
        price.extend(df['PRICE'].dropna().astype(int).tolist())
    
    # Calculate total cards
    #total_cards.extend((df['Qty Ordered'] * df['CARDS PER UNIT'].dropna()).tolist())
    # Get total cards
    arr1 = sell_units[1:]
    arr2 = cards_per_units[1:]
    for a,b in zip(arr1,arr2):
        num = int(a)*int(b)
        total_cards.append(num)
    

    # Titles of columns
    output_data = (
        delivery_dates,
        ship_to_name,
        ship_to_address,
        ship_to_city,
        ship_to_state,
        ship_to_zip,
        ship_to_contact,
        po_nums,
        sell_units,
        cards_per_units,
        total_cards,
    )
    return output_data

def get_sku_num(file_path):
    _, file_extension = os.path.splitext(file_path)
    if file_extension == ".numbers":
        # Load the .numbers file
        doc = Document(file_path)
        sheets = doc.sheets
        tables = sheets[0].tables
        rows = tables[0].rows()

        # Scan input file
        header_row = rows[0]
        sku_num.extend(get_column_data(rows, header_row, "Buyers Catalog or Stock Keeping #"))
    elif file_extension == ".csv":
        # Load the CSV file into a DataFrame
        df = pd.read_csv(file_path)
        # Get data from the DataFrame
        sku_num.extend(get_csv_column_info(df, "Buyers Catalog or Stock Keeping #"))
            
    else:
        print(f"{file_path} is neither a Numbers file nor a CSV file.")
    return sku_num

def convert_numbers_data(file_name):
    # Load the .numbers file
    doc = Document(file_name)
    sheets = doc.sheets
    tables = sheets[0].tables
    rows = tables[0].rows()

    # Scan input file
    header_row = rows[0]
    print(header_row)

    # Get PO Numbers using column NAME
    column_name = "PO Number"
    column_index = None
    for idx, cell in enumerate(header_row):
        if cell.value == column_name:
            column_index = idx
            break

    for row in rows[1:]:
        cell = row[column_index]
        if cell.value is not None and cell.value not in po_nums:
            po_nums.append(cell.value)
    for value in po_nums:
        print(value)
        
    # Get Ship To Name
    ship_to_name.extend(get_column_data(rows, header_row, "Ship To Name"))
        
    # Get Ship To Address 1
    ship_to_address.extend(get_column_data(rows, header_row, "Ship To Address 1"))
        
    # Get Ship To State
    ship_to_state.extend(get_column_data(rows, header_row, "Ship To State"))
        
    # Get Ship To Zip
    ship_to_zip.extend(get_column_data(rows, header_row, "Ship to Zip"))
        
    # Get Ship To City
    ship_to_city.extend(get_column_data(rows, header_row, "Ship To City"))
        
    # Get Ship To Contact
    ship_to_contact.extend(get_column_data(rows, header_row, "Ship To Contact"))
    
    sku_num.extend(get_column_data(rows, header_row, "Buyers Catalog or Stock Keeping #"))

    # Get sell units
    column_name = "Qty Ordered"
    column_index = None
    for idx, cell in enumerate(header_row):
        if cell.value == column_name:
            column_index = idx
            break
    for row in rows[1:]:
        cell = row[column_index]
        if cell.value is not None:
            sell_units.append(cell.value)
    for value in sell_units:
        print(value)
        
    # Get cards per Order
    column_name = "Product/Item Description"
    column_index = None
    for idx, cell in enumerate(header_row):
        if cell.value == column_name:
            column_index = idx
            break
    for row in rows[1:]:
        cell = row[column_index]
        if cell.value is not None:
            string_val = cell.value
            match = re.search(r'(\d+)/\$', string_val) # Find the number before the '/$' symbol
            if match:
                number = match.group(1)
                cards_per_units.append(number)
            match = re.search(r'\$(\d+(?:\.\d+)?)', string_val) # Find the number before the '/$' symbol
            if match:
                number = match.group(1)
                price.append(number)

    # Get total cards
    arr1 = sell_units[1:]
    arr2 = cards_per_units[1:]
    for a,b in zip(arr1,arr2):
        num = int(a)*int(b)
        total_cards.append(num)

    # Get delivery dates
    delivery_dates.extend(get_column_data(rows, header_row, "Requested Delivery Date"))
    
    # Titles of columns
    output_data = (
        delivery_dates,
        ship_to_name,
        ship_to_address,
        ship_to_city,
        ship_to_state,
        ship_to_zip,
        ship_to_contact,
        po_nums,
        sell_units,
        cards_per_units,
        total_cards,
    )
    return output_data

def replace_word_in_docx(file_path, old_word, new_word, new_file_name):
    # Load the Word document
    doc = WordDocument(file_path)
    
    
    # Iterate through each paragraph
    for paragraph in doc.paragraphs:
        if old_word in paragraph.text:
            # Replace the old word with the new word
            paragraph.text = paragraph.text.replace(old_word, new_word)
            
    
    # Save the modified document as a new file
    doc.save(new_file_name)
    print(f"Replaced '{old_word}' with '{new_word}' and saved as {new_file_name}")




# Serve the static files from the React build
@app.route('/')
@cross_origin()
def serve():
    return send_from_directory(app.static_folder, 'index.html')

@app.route('/upload', methods=['POST'])
@cross_origin()
def upload_file():
    global delivery_dates, ship_to_address, ship_to_name, ship_to_city, ship_to_state, ship_to_contact, ship_to_zip, po_nums, total_cards, sell_units,cards_per_units, price, sku_num
    # Variables for the output file
    delivery_dates = ['Delivery Date']
    ship_to_name = ['Ship to Name']
    ship_to_address = ['Ship to Address']
    ship_to_city = ['Ship to City']
    ship_to_state = ['Ship to State']
    ship_to_zip = ['Ship to Zip']
    ship_to_contact = ['Ship to Contact']
    po_nums = ['PO Number']
    sell_units = ['SELL UNITS']
    cards_per_units = ['CARDS PER UNIT']
    price = ['PRICE']
    total_cards = ['TOTAL CARDS']
    sku_num = ['SKU Number']
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400

    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400

    # Process the File
    # Process file (for example, modify a CSV file)
    if file.filename.endswith('.xlsx'):
        return jsonify({'message': 'Excel file processed successfully'}), 200
    elif file.filename.endswith('.csv'):
        # Create file
        output_workbook = xlsxwriter.Workbook(processed_file_path)
        worksheet = output_workbook.add_worksheet()
        output_data = convert_csv_data(file)
        

        # Iterate over data and write to columns
        col = 0
        row = 0
        for names in (output_data):
            for x in names:
                if col == 7: 
                    bold_format = output_workbook.add_format({'bold': True})
                    worksheet.write(row, col, x, bold_format)
                else:
                    worksheet.write(row, col, x)
                row +=1
            col +=1
            row = 0
            
        header_row_format = output_workbook.add_format({'bold': True})
        worksheet.set_row(0, None, header_row_format)
        # Set the width of the columns to fit the content
        worksheet.set_column('A:A', 15)  # Column A: Delivery Dates
        worksheet.set_column('B:B', 20)  # Column B: Ship To Name
        worksheet.set_column('C:C', 15)  # Column C: Ship To Address
        worksheet.set_column('D:D', 15)  # Column D: Ship To City
        worksheet.set_column('E:E', 10)  # Column E: Ship To State
        worksheet.set_column('F:F', 10)  # Column F: Ship To Zip
        worksheet.set_column('G:G', 15)  # Column G: Ship To Contact
        worksheet.set_column('H:H', 20)  # Column H: PO Numbers
        worksheet.set_column('I:I', 15)  # Column I: Sell Units
        worksheet.set_column('J:J', 15)  # Column J: Cards Per Unit
        worksheet.set_column('K:K', 15)  # Column K: Total Cards
        #sku_num.extend(get_sku_num("/Users/siddiepennewill/Desktop/test-csvs/Bear's Smokehouse - po014930816274_13_2720.00_337343.csv"))
        replace1 = "	PO-NUM	  (LOCATION)	 	"
        replace2= "   	SKU-NUM	"
        replace3 = "    XX	"
        replace4 = "    	YY  $ZZ Gift Cards		"
        new_word1 = "	"+str(po_nums[1])+"	  ("+str(ship_to_city[1])+")	 	"
        new_word2 = "   	"+str(sku_num[1])+"	"
        new_word3= "    "+str(sell_units[1])+"	"
        new_word4="    	"+str(cards_per_units[1])+"  $"+str(price[1])+" Gift Cards		"
        replace_word_in_docx("./HAND-DELIVERY-FORM.docx", replace1, new_word1, "./HandDeliveryForm.docx")
        replace_word_in_docx("./HandDeliveryForm.docx", replace2, new_word2, "./HandDeliveryForm.docx")
        replace_word_in_docx("./HandDeliveryForm.docx", replace3, new_word3, "./HandDeliveryForm.docx")
        replace_word_in_docx("./HandDeliveryForm.docx", replace4, new_word4, "./HandDeliveryForm.docx")
  
        output_workbook.close()
        
        return jsonify({
            'message': 'File processed successfully',
            'ship_to_name': ship_to_name,
            'delivery_dates': delivery_dates,
            'units': sell_units
        }), 200
    
    elif file.filename.endswith('.numbers'):
        # Save the file temporarily
        temp_file_path = os.path.join('/tmp', file.filename)
        file.save(temp_file_path)
        
        # Create file
        output_workbook = xlsxwriter.Workbook('COSTCO.xlsx')
        worksheet = output_workbook.add_worksheet()
        output_data = convert_numbers_data(temp_file_path)
        
        # Iterate over data and write to columns
        col = 0
        row = 0
        for names in (output_data):
            for x in names:
                if col == 7: 
                    bold_format = output_workbook.add_format({'bold': True})
                    worksheet.write(row, col, x, bold_format)
                else:
                    worksheet.write(row, col, x)
                row +=1
            col +=1
            row = 0
            
        header_row_format = output_workbook.add_format({'bold': True})
        worksheet.set_row(0, None, header_row_format)
        # Set the width of the columns to fit the content
        worksheet.set_column('A:A', 15)  # Column A: Delivery Dates
        worksheet.set_column('B:B', 20)  # Column B: Ship To Name
        worksheet.set_column('C:C', 15)  # Column C: Ship To Address
        worksheet.set_column('D:D', 15)  # Column D: Ship To City
        worksheet.set_column('E:E', 10)  # Column E: Ship To State
        worksheet.set_column('F:F', 10)  # Column F: Ship To Zip
        worksheet.set_column('G:G', 15)  # Column G: Ship To Contact
        worksheet.set_column('H:H', 20)  # Column H: PO Numbers
        worksheet.set_column('I:I', 15)  # Column I: Sell Units
        worksheet.set_column('J:J', 15)  # Column J: Cards Per Unit
        worksheet.set_column('K:K', 15)  # Column K: Total Cards
        #sku_num.extend(get_sku_num("/Users/siddiepennewill/Desktop/test-csvs/Bear's Smokehouse - po014930816274_13_2720.00_337343.csv"))
  
        replace1 = "	PO-NUM	  (LOCATION)	 	"
        replace2= "   	SKU-NUM	"
        replace3 = "    XX	"
        replace4 = "    	YY  $ZZ Gift Cards		"
        new_word1 = "	"+str(po_nums[1])+"	  ("+str(ship_to_city[1])+")	 	"
        new_word2 = "   	"+str(sku_num[1])+"	"
        new_word3= "    "+str(sell_units[1])+"	"
        new_word4="    	"+str(cards_per_units[1])+"  $"+str(price[1])+" Gift Cards		"
        replace_word_in_docx("./HAND-DELIVERY-FORM.docx", replace1, new_word1, "./HandDeliveryForm.docx")
        replace_word_in_docx("./HandDeliveryForm.docx", replace2, new_word2, "./HandDeliveryForm.docx")
        replace_word_in_docx("./HandDeliveryForm.docx", replace3, new_word3, "./HandDeliveryForm.docx")
        replace_word_in_docx("./HandDeliveryForm.docx", replace4, new_word4, "./HandDeliveryForm.docx")
        output_workbook.close()
        
        return jsonify({
            'message': 'File processed successfully',
            'ship_to_name': ship_to_name,
            'delivery_dates': delivery_dates,
            'units': sell_units
        }), 200
        
    else: 
        return jsonify({'error': 'Unsupported file type'}), 400
    
        
@app.route('/download', methods=['GET'])
@cross_origin()
def download_file():
    # Check if the file exists
    if os.path.exists(processed_file_path):
        return send_file(
            processed_file_path,
            as_attachment=True,
            download_name='new-file.xlsx',
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    else:
        return jsonify({'error': 'File not found'}), 404

@app.route('/generate-docx', methods=['POST'])
@cross_origin()
def generate_docx():
    file_path_docx = "./HandDeliveryForm.docx"
    if os.path.exists(file_path_docx):
        return send_file(
            file_path_docx,
            as_attachment=True,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
    else:
        return jsonify({'error': 'File not found'}), 404
    


if __name__ == '__main__':
    app.run(debug=True)
